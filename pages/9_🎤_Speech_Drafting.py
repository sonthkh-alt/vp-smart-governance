import streamlit as st
import os
import sys
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from utils.ai_helper import generate_document_content
from utils.doc_helper import extract_text_from_pdf, extract_text_from_docx
from utils.ui_helper import set_premium_css, draw_module_header, draw_sidebar
from utils.auth_helper import require_auth
import database
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Khởi tạo DB
database.init_db()

# Áp dụng giao diện Premium
set_premium_css()

# Hiển thị Header
draw_module_header(
    "Speech Drafting",
    "🎤",
    "Trợ lý soạn thảo bài phát biểu chuyên nghiệp dành cho Lãnh đạo Đoàn ĐBQH & HĐND."
)

# Khởi tạo session state
if "speech_draft" not in st.session_state:
    st.session_state.speech_draft = ""
if "speech_docx_bytes" not in st.session_state:
    st.session_state.speech_docx_bytes = None

# -- SIDEBAR: LỊCH SỬ --
with st.sidebar:
    st.markdown("### 📜 Lịch sử Phát biểu")
    drafts = database.get_drafts()
    speech_drafts = [d for d in drafts if d['doc_type'] == "Bài phát biểu"]
    
    if speech_drafts:
        for d in speech_drafts[:5]:
            if st.button(f"🎤 {d['created_at']}", key=f"speech_{d['id']}", use_container_width=True):
                st.session_state.speech_draft = d['ai_content'].get('noi_dung_chinh', '')
                st.success("Đã tải lại bài phát biểu cũ!")
    else:
        st.caption("Chưa có lịch sử soạn thảo bài phát biểu.")

# -- MAIN DASHBOARD --
col1, col2 = st.columns([1, 1], gap="large")

with col1:
    st.markdown("### 📝 1. Thông tin Yêu cầu")
    
    speech_type = st.selectbox(
        "Loại bài phát biểu:",
        [
            "Khai mạc Kỳ họp HĐND", 
            "Bế mạc Kỳ họp HĐND", 
            "Phát biểu chỉ đạo tại Hội nghị", 
            "Bài phát biểu chúc mừng (Lễ, Tết)", 
            "Phát biểu tại buổi Tiếp xúc Cử tri", 
            "Khác"
        ]
    )
    
    speaker_role = st.text_input("Chức danh người phát biểu:", placeholder="Ví dụ: Chủ tịch HĐND tỉnh Thanh Hóa")
    event_name = st.text_input("Tên sự kiện/buổi lễ:", placeholder="Ví dụ: Kỳ họp thứ 20, HĐND tỉnh khóa XVIII")
    
    key_points = st.text_area(
        "Các ý chính cần nhấn mạnh:",
        height=150,
        placeholder="Ví dụ: \n- Đánh giá cao kết quả kinh tế - xã hội quý I.\n- Nhấn mạnh nhiệm vụ giải ngân vốn đầu tư công.\n- Kêu gọi tinh thần đoàn kết, đổi mới..."
    )
    
    st.markdown("#### 📁 Tài liệu tham khảo (Dữ liệu gốc)")
    ref_files = st.file_uploader("Tải lên báo cáo, nghị quyết hoặc dữ liệu liên quan:", type=["pdf", "docx"], accept_multiple_files=True)
    
    if st.button("🚀 SOẠN THẢO BÀI PHÁT BIỂU", use_container_width=True, type="primary"):
        if require_auth("Soạn thảo bài phát biểu"):
            if not key_points.strip() and not ref_files:
                st.warning("Vui lòng nhập các ý chính hoặc tải lên tài liệu tham khảo!")
            else:
                with st.spinner("AI đang nghiên cứu ngữ cảnh và soạn thảo bài phát biểu..."):
                    # Đọc ngữ cảnh
                    context_text = f"Người phát biểu: {speaker_role}\nSự kiện: {event_name}\n"
                    if ref_files:
                        for f in ref_files:
                            if f.name.endswith(".pdf"):
                                context_text += extract_text_from_pdf(f) + "\n"
                            else:
                                context_text += extract_text_from_docx(f) + "\n"
                    
                    prompt = f"Hãy soạn thảo bài phát biểu chuyên nghiệp cho {speaker_role} tại {event_name}. Các nội dung trọng tâm cần có: {key_points}. Văn phong phải hùng hồn, truyền cảm hứng, trang trọng nhưng gần gũi với nhân dân."
                    
                    result = generate_document_content(prompt, doc_type="Bài phát biểu", context=context_text)
                    
                    if isinstance(result, dict) and "error" not in result:
                        st.session_state.speech_draft = result.get("noi_dung_chinh", "")
                        database.save_draft("Bài phát biểu", prompt, result)
                        
                        user_email = st.session_state.user_info.get("email") if st.session_state.get("is_logged_in") else "Khách"
                        database.log_action(user_email, "Soạn thảo bài phát biểu", "Phát biểu", f"Loại: {speech_type}")
                        st.success("Đã soạn thảo xong bài phát biểu!")
                    else:
                        st.error("Có lỗi xảy ra khi tạo nội dung.")

with col2:
    st.markdown("### ✨ 2. Nội dung Bài phát biểu")
    
    if st.session_state.speech_draft:
        edited_speech = st.text_area(
            "Bạn có thể chỉnh sửa nội dung bên dưới:",
            value=st.session_state.speech_draft,
            height=450
        )
        st.session_state.speech_draft = edited_speech
        
        st.markdown("---")
        
        if st.button("📄 Xuất file Word chuyên nghiệp", use_container_width=True, type="primary"):
            with st.spinner("Đang tạo file Word..."):
                doc = Document()
                # Cấu hình font
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(14)
                
                # Tiêu đề bài phát biểu
                title = doc.add_paragraph()
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = title.add_run(f"BÀI PHÁT BIỂU CỦA {speaker_role.upper()}\nTẠI {event_name.upper()}")
                run.bold = True
                run.font.size = Pt(16)
                
                doc.add_paragraph() # Khoảng trống
                
                # Nội dung
                for line in st.session_state.speech_draft.split('\n'):
                    p = doc.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Footer
                doc.add_paragraph()
                footer = doc.add_paragraph("Thanh Hóa, ngày ... tháng ... năm ...")
                footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                output = io.BytesIO()
                doc.save(output)
                st.session_state.speech_docx_bytes = output.getvalue()
                st.success("Đã tạo file thành công!")
        
        if st.session_state.speech_docx_bytes:
            st.download_button(
                label="⬇️ TẢI XUỐNG BÀI PHÁT BIỂU (.DOCX)",
                data=st.session_state.speech_docx_bytes,
                file_name=f"BaiPhatBieu_{speech_type.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )
    else:
        st.info("Nội dung bài phát biểu sẽ hiển thị tại đây sau khi bạn hoàn tất bước 1.")

st.markdown("---")
st.caption("Giao diện Soạn thảo Bài phát biểu - Hỗ trợ AI tối ưu cho Lãnh đạo HĐND.")
