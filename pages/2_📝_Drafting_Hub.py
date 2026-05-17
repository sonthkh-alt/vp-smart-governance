import os
os.environ["PROTOCOL_BUFFERS_PYTHON_IMPLEMENTATION"] = "python"

import streamlit as st
import sys
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from utils.ai_helper import generate_document_content, check_legal_compliance
from utils.doc_helper import extract_text_from_pdf, extract_text_from_docx, create_nd30_document, create_tham_tra_document
from utils.ui_helper import set_premium_css, draw_module_header
from utils.auth_helper import require_auth
from utils.gemini_client import generate_text
import database
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Khởi tạo DB
database.init_db()
set_premium_css()

# Khởi tạo Session State
if "rev_text_content" not in st.session_state:
    st.session_state.rev_text_content = ""

draw_module_header(
    "Trung tâm Soạn thảo & Kiểm soát Chất lượng",
    "📝",
    "Hệ sinh thái AI hỗ trợ soạn thảo văn bản hành chính, bài phát biểu và soát xét lỗi chuyên nghiệp."
)
st.caption("🚀 Phiên bản: 2.5.1 (Cập nhật tính năng Upload file)")

tab1, tab2, tab3, tab4 = st.tabs(["📝 Soạn thảo Văn bản", "🎤 Soạn thảo Phát biểu", "📋 Soát xét & Kiểm lỗi", "🔍 Thẩm tra Dự thảo NQ"])

# --- Cấu hình AI chung cho trang ---
with st.sidebar:
    st.markdown("### 🤖 Cấu hình AI")
    ai_provider = st.radio("Chọn mô hình AI ưu tiên:", ["Anthropic Claude (Mặc định)", "OpenAI ChatGPT", "Google Gemini", "Groq (Llama 3.3)"], index=0, key="global_ai_provider")
    if "Anthropic" in ai_provider:
        provider_key = "claude"
    elif "OpenAI" in ai_provider:
        provider_key = "openai"
    elif "Gemini" in ai_provider:
        provider_key = "gemini"
    else:
        provider_key = "groq"

# --- TAB 1: SOẠN THẢO VĂN BẢN ---
with tab1:
    st.markdown("### 📝 Tham mưu & Soạn thảo chuẩn NĐ 30")
    col1, col2 = st.columns([1, 1], gap="large")
    with col1:
        doc_type = st.selectbox("Loại văn bản:", ["Tự động", "Công văn", "Báo cáo", "Quyết định", "Nghị quyết", "Tờ trình"], key="draft_type")
        prompt = st.text_area("Yêu cầu chi tiết:", height=150, placeholder="Ví dụ: Soạn công văn trả lời đề xuất kinh phí...", key="draft_prompt")
        legal_data = st.text_area("Căn cứ pháp lý (NotebookLM):", height=150, placeholder="Dán nội dung luật hoặc quy định tại đây...", key="draft_legal")
        ref_files = st.file_uploader("Tài liệu tham khảo:", type=["pdf", "docx"], accept_multiple_files=True, key="draft_ref")
        # Double column layout for buttons
        c1_btn, c2_btn = st.columns([2, 1])
        with c1_btn:
            if st.button("🚀 TẠO DỰ THẢO VĂN BẢN", type="primary", use_container_width=True, key="draft_run"):
                if require_auth("Soạn thảo văn bản"):
                    with st.spinner(f"AI ({ai_provider}) đang soạn thảo..."):
                        context = ""
                        if ref_files:
                            for f in ref_files: context += (extract_text_from_pdf(f) if f.name.endswith(".pdf") else extract_text_from_docx(f)) + "\n"
                        res = generate_document_content(prompt, doc_type=doc_type, context=context, notebook_lm_data=legal_data, provider=provider_key)
                        if isinstance(res, dict) and "error" not in res:
                            st.session_state.draft_res = res.get("noi_dung_chinh", "")
                            st.session_state.draft_edit = res.get("noi_dung_chinh", "")
                            st.session_state.draft_dict = res
                            database.save_draft(doc_type, prompt, res)
                        else: st.error(f"Lỗi: {res.get('error') if isinstance(res, dict) else res}")
        with c2_btn:
            if st.button("🧹 LÀM MỚI", use_container_width=True, key="draft_clear"):
                for k in ["draft_res", "draft_edit", "draft_dict", "draft_prompt", "draft_legal"]:
                    st.session_state.pop(k, None)
                st.rerun()
    with col2:
        if "draft_res" in st.session_state:
            edited = st.text_area("Nội dung dự thảo:", value=st.session_state.draft_res, height=400, key="draft_edit")
            st.session_state.draft_res = edited
            if st.button("📄 Xuất file Word chuẩn NĐ 30", use_container_width=True, type="primary", key="draft_word"):
                dict_meta = st.session_state.get("draft_dict", {})
                dict_meta["noi_dung_chinh"] = edited
                out = create_nd30_document(dict_meta)
                st.download_button("⬇️ Tải xuống (.docx)", out.getvalue(), "VanBan_DuThao.docx", key="draft_dl")
        else: st.info("Dự thảo sẽ hiển thị ở đây.")

# --- TAB 2: SOẠN THẢO PHÁT BIỂU ---
with tab2:
    st.markdown("### 🎤 Trợ lý soạn thảo bài phát biểu Lãnh đạo")
    c1, c2 = st.columns([1, 1], gap="large")
    with c1:
        s_role = st.text_input("Chức danh người phát biểu:", placeholder="Ví dụ: Chủ tịch HĐND tỉnh", key="speech_role")
        s_event = st.text_input("Tên sự kiện:", placeholder="Ví dụ: Kỳ họp thứ 20 HĐND tỉnh", key="speech_event")
        s_key = st.text_area("Ý chính cần nhấn mạnh:", height=150, key="speech_key")
        # Double column layout for speech buttons
        sp_c1, sp_c2 = st.columns([2, 1])
        with sp_c1:
            if st.button("🚀 SOẠN THẢO BÀI PHÁT BIỂU", type="primary", use_container_width=True, key="speech_run"):
                if require_auth("Soạn thảo phát biểu"):
                    prompt = f"Soạn bài phát biểu cho {s_role} tại {s_event}. Ý chính: {s_key}"
                    with st.spinner(f"AI ({ai_provider}) đang soạn thảo..."):
                        res = generate_document_content(prompt, doc_type="Bài phát biểu", provider=provider_key)
                        if isinstance(res, dict) and "error" not in res:
                            st.session_state.speech_res = res.get("noi_dung_chinh", "")
                            st.session_state.speech_edit = res.get("noi_dung_chinh", "")
                            database.save_draft("Bài phát biểu", prompt, res)
                        else: st.error("Lỗi tạo nội dung.")
        with sp_c2:
            if st.button("🧹 LÀM MỚI", use_container_width=True, key="speech_clear"):
                for k in ["speech_res", "speech_edit", "speech_role", "speech_event", "speech_key"]:
                    st.session_state.pop(k, None)
                st.rerun()
    with c2:
        if "speech_res" in st.session_state:
            s_edit = st.text_area("Nội dung bài phát biểu:", value=st.session_state.speech_res, height=400, key="speech_edit")
            st.session_state.speech_res = s_edit
            if st.button("📄 Xuất file Word Phát biểu", use_container_width=True, key="speech_word"):
                doc = Document()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"BÀI PHÁT BIỂU CỦA {s_role.upper()}\nTẠI {s_event.upper()}")
                run.bold = True
                for line in s_edit.split('\n'): doc.add_paragraph(line)
                out = io.BytesIO(); doc.save(out)
                st.download_button("⬇️ Tải xuống (.docx)", out.getvalue(), "BaiPhatBieu.docx", key="speech_dl")
        else: st.info("Bài phát biểu sẽ hiển thị ở đây.")

# --- TAB 3: SOÁT XÉT & KIỂM LỖI ---
with tab3:
    st.markdown("### 📋 Kiểm soát Chất lượng & Văn phong")
    cl, cr = st.columns([1, 1.3], gap="large")
    with cl:
        st.markdown("#### 📥 Nhập Văn bản")
        rev_file = st.file_uploader("Tải lên văn bản (PDF/DOCX):", type=["pdf", "docx"], key="rev_file")
        
        # Xử lý trích xuất text từ file upload
        if rev_file and ("last_rev_file" not in st.session_state or st.session_state.last_rev_file != rev_file.name):
            with st.spinner("🔍 Đang trích xuất văn bản..."):
                extracted_text = ""
                if rev_file.name.endswith(".pdf"):
                    extracted_text = extract_text_from_pdf(rev_file)
                else:
                    extracted_text = extract_text_from_docx(rev_file)
                
                if extracted_text.strip():
                    st.session_state.rev_text_content = extracted_text
                else:
                    st.warning("⚠️ Không thể trích xuất văn bản từ file này hoặc file trống.")
                st.session_state.last_rev_file = rev_file.name

        # Hiển thị text area - Tự động đồng bộ với st.session_state.rev_text_content
        st.text_area("Nội dung văn bản cần kiểm tra:", height=300, key="rev_text_content")
        
        rev_focus = st.multiselect("Trọng tâm kiểm tra:", ["Chính tả & Ngữ pháp", "Thể thức NĐ 30", "Văn phong hành chính", "Logic quản lý"], default=["Chính tả & Ngữ pháp", "Văn phong hành chính"], key="rev_focus")
        # Double column layout for audit buttons
        rev_c1, rev_c2 = st.columns([2, 1])
        with rev_c1:
            if st.button("🚀 BẮT ĐẦU KIỂM TRA", type="primary", use_container_width=True, key="rev_run"):
                if require_auth("Kiểm tra văn bản"):
                    if not st.session_state.rev_text_content.strip():
                        st.error("⚠️ Vui lòng tải file hoặc nhập nội dung văn bản!")
                    else:
                        with st.spinner(f"AI ({ai_provider}) đang soát lỗi..."):
                            p = f"""
Bạn là Chuyên gia Kiểm soát Chất lượng & Pháp chế Hành chính cao cấp của Văn phòng Đoàn ĐBQH và HĐND tỉnh Thanh Hóa.
Nhiệm vụ của bạn là rà soát toàn bộ văn bản dưới đây, phát hiện ra TẤT CẢ các lỗi sai từ nhỏ nhất đến lớn nhất và đề xuất sửa đổi cụ thể.

YÊU CẦU BẮT BUỘC:
1. CHỈ RA LỖI SAI & ĐỀ XUẤT SỬA: Không viết lại toàn bộ văn bản gốc. Chỉ liệt kê chi tiết các lỗi phát hiện được vào bảng.
2. PHÁT HIỆN LỖI SƠ ĐẲNG (CỰC KỲ QUAN TRỌNG): Hãy rà soát kỹ từng từ một để phát hiện các lỗi gõ phím bừa, từ vô nghĩa (Ví dụ: "fádf", "asdf", "ghjk"), từ viết sai chính tả tiếng Việt, ký tự lạ/ký tự rác hoặc từ bị dính chữ. Phải chỉ ra chính xác vị trí và yêu cầu xóa bỏ hoặc sửa lại.
3. THỂ THỨC NĐ 30 & VĂN PHONG HÀNH CHÍNH (VĂN PHÒNG): Phát hiện các lỗi viết hoa tùy tiện, dùng từ không chuẩn ngôn ngữ công vụ, cách đặt tiêu đề báo cáo không khớp với nội dung, hoặc thẩm quyền ký/nơi nhận không hợp lý.
4. LỖI LOGIC QUẢN LÝ & SỰ NHẤT QUÁN (ĐẶC BIỆT SÂU SẮC): 
   - Kiểm tra tính mâu thuẫn trong số liệu thống kê ở các phần khác nhau của tài liệu (Ví dụ: tổng số xã miền núi/đồng bằng cộng lại không bằng tổng số chung).
   - Kiểm tra tính hợp lý của niên độ báo cáo và mốc thời gian (Ví dụ: báo cáo sơ kết "01 năm" nhưng khoảng thời gian thống kê từ 01/7/2025 đến 15/5/2026 là chưa tròn 12 tháng - đây là lỗi logic quản lý rất nghiêm trọng!).
   - Phát hiện các điểm đứt gãy, thiếu mạch lạc logic giữa phần đánh giá thực trạng khó khăn và phần đề xuất giải pháp, phương hướng (Đề xuất không giải quyết được khó khăn nêu ra).

TRỌNG TÂM KIỂM TRA ĐƯỢC YÊU CẦU: {', '.join(rev_focus)}

HÃY TRÌNH BÀY BÁO CÁO SOÁT LỖI THEO CẤU TRÚC PHÂN TÍCH SAU:

### 📊 1. BẢNG PHÂN TÍCH CHI TIẾT CÁC LỖI PHÁT HIỆN & ĐỀ XUẤT HIỆU CHỈNH
| STT | Vị trí (Dòng/Đoạn) | Nội dung gốc (Có lỗi) | Đề xuất hiệu chỉnh | Phân loại lỗi | Chi tiết lý do & Căn cứ sửa đổi |
|---|---|---|---|---|---|
| 1 | [Ví dụ: Đoạn 1] | [Nội dung lỗi] | [Đề xuất sửa] | [Chính tả / Thể thức NĐ 30 / Văn phong] | [Lý do & Căn cứ quy định] |
| ... | ... | ... | ... | ... | ... |

---

### 📈 2. ĐÁNH GIÁ CHẤT LƯỢNG VĂN BẢN CHUNG
* **Điểm đánh giá chất lượng hiện tại:** [Cho điểm trên thang 10]
* **Tổng hợp các vấn đề chính:** [Tóm tắt ngắn gọn các lỗi hệ thống hoặc lỗi lặp đi lặp lại nhiều nhất]
* **Khuyến nghị cải thiện:** [Các hướng dẫn chuyên gia để tác giả tự hoàn thiện văn bản tốt hơn]

Nội dung văn bản cần soát lỗi:
{st.session_state.rev_text_content}
"""
                            res = generate_text(p, use_pro=True, provider=provider_key, use_search=False)
                            st.session_state.rev_res = res
        with rev_c2:
            if st.button("🧹 LÀM MỚI", use_container_width=True, key="rev_clear"):
                for k in ["rev_res", "rev_text_content", "last_rev_file"]:
                    st.session_state.pop(k, None)
                st.rerun()
    with cr:
        if "rev_res" in st.session_state:
            st.markdown(st.session_state.rev_res)
            st.download_button("⬇️ Tải Báo cáo (.txt)", st.session_state.rev_res.encode("utf-8"), "BaoCao_SoatLoi.txt", key="rev_dl")
        else: st.info("Kết quả soát lỗi sẽ hiển thị ở đây.")

# --- TAB 4: THẨM TRA DỰ THẢO NGHỊ QUYẾT ---
with tab4:
    st.markdown("### 🔍 Thẩm tra Dự thảo Nghị quyết HĐND tỉnh")
    st.caption("Hệ thống AI chuyên gia thẩm tra tự động phân tích Tờ trình, Dự thảo NQ và các văn bản liên quan để tạo Báo cáo thẩm tra chuyên nghiệp.")
    
    tt_col1, tt_col2 = st.columns([1, 1.2], gap="large")
    with tt_col1:
        st.markdown("#### 📥 Tài liệu đầu vào")
        ban_select = st.selectbox("🏛️ Ban thẩm tra:", [
            "Kinh tế - Ngân sách",
            "Pháp chế",
            "Văn hóa - Xã hội",
            "Dân tộc"
        ], key="tt_ban")
        
        nq_name = st.text_input("📋 Tên Dự thảo Nghị quyết:", placeholder="Ví dụ: Nghị quyết về phân bổ ngân sách địa phương năm 2026", key="tt_nq_name")
        
        tt_to_trinh = st.file_uploader("📄 Tờ trình của UBND tỉnh:", type=["pdf", "docx"], key="tt_totrinh")
        tt_du_thao = st.file_uploader("📄 Dự thảo Nghị quyết:", type=["pdf", "docx"], key="tt_duthao")
        tt_lien_quan = st.file_uploader("📄 Văn bản liên quan (nếu có):", type=["pdf", "docx"], accept_multiple_files=True, key="tt_lienquan")
        
        tt_note = st.text_area("📝 Ghi chú / Yêu cầu đặc biệt (nếu có):", height=100, placeholder="Ví dụ: Tập trung phản biện nguồn kinh phí bố trí...", key="tt_note")
        
        tt_c1, tt_c2 = st.columns([2, 1])
        with tt_c1:
            if st.button("🚀 TẠO BÁO CÁO THẨM TRA", type="primary", use_container_width=True, key="tt_run"):
                if require_auth("Thẩm tra dự thảo NQ"):
                    if not tt_to_trinh and not tt_du_thao:
                        st.error("⚠️ Vui lòng tải lên ít nhất Tờ trình hoặc Dự thảo Nghị quyết!")
                    else:
                        with st.spinner(f"AI ({ai_provider}) đang thẩm tra chuyên sâu..."):
                            # Trích xuất nội dung các tài liệu
                            to_trinh_text = ""
                            du_thao_text = ""
                            lien_quan_text = ""
                            
                            if tt_to_trinh:
                                to_trinh_text = extract_text_from_pdf(tt_to_trinh) if tt_to_trinh.name.endswith(".pdf") else extract_text_from_docx(tt_to_trinh)
                            if tt_du_thao:
                                du_thao_text = extract_text_from_pdf(tt_du_thao) if tt_du_thao.name.endswith(".pdf") else extract_text_from_docx(tt_du_thao)
                            if tt_lien_quan:
                                for f in tt_lien_quan:
                                    lien_quan_text += (extract_text_from_pdf(f) if f.name.endswith(".pdf") else extract_text_from_docx(f)) + "\n---\n"
                            
                            prompt = f"""Bạn là Chuyên gia thẩm tra cao cấp của Ban {ban_select} - Hội đồng nhân dân tỉnh Thanh Hóa.
Bạn đang thực hiện thẩm tra dự thảo Nghị quyết trước khi trình Kỳ họp HĐND tỉnh.

NHIỆM VỤ THẨM TRA:
Viết BÁO CÁO THẨM TRA hoàn chỉnh, chuyên nghiệp theo đúng cấu trúc của báo cáo thẩm tra chuẩn HĐND, gồm 4 phần chính:

I. CƠ SỞ PHÁP LÝ VÀ THẨM QUYỀN BAN HÀNH
- Căn cứ pháp lý: Liệt kê và đánh giá các căn cứ luật, pháp lệnh, nghị định mà Tờ trình viện dẫn.
- Thẩm quyền ban hành: Xác nhận HĐND tỉnh có thẩm quyền ban hành nghị quyết này không.
- Trình tự thủ tục: Đánh giá quy trình xây dựng dự thảo có đúng quy định không.

II. SỰ PHÙ HỢP VỀ NỘI DUNG
- Phân tích sự phù hợp của nội dung dự thảo với chủ trương, đường lối của Đảng, chính sách pháp luật của Nhà nước.
- Đánh giá tính khả thi, phù hợp với điều kiện thực tiễn địa phương.
- Phát hiện các điểm bất cập, thiếu sót, mâu thuẫn trong nội dung dự thảo.
- So sánh đối chiếu giữa Tờ trình và Dự thảo NQ: nội dung có nhất quán không.

III. KHẢ NĂNG CÂN ĐỐI NGUỒN LỰC
- Phản biện nguồn kinh phí: Đánh giá tính hợp lý của dự toán kinh phí.
- Nguồn nhân lực: Đánh giá năng lực tổ chức thực hiện.
- Khả năng cân đối ngân sách địa phương.
- Rủi ro tài chính và các phương án dự phòng.

IV. KẾT LUẬN VÀ KIẾN NGHỊ
- Ý kiến thống nhất hoặc không thống nhất với nội dung dự thảo.
- Các nội dung đề nghị chỉnh sửa, bổ sung cụ thể.
- Kiến nghị HĐND tỉnh xem xét, quyết định.

YÊU CẦU VĂN PHONG:
- Ngôn ngữ trang trọng, chuẩn mực công vụ, văn phong thẩm tra chuyên nghiệp.
- Phân tích sâu sắc, có lập luận chặt chẽ, dẫn chứng cụ thể từ tài liệu.
- KHÔNG dùng markdown heading (#, ##), chỉ dùng ký hiệu I., II., 1., 2., a), b) theo chuẩn văn bản hành chính.

{f"GHI CHÚ CỦA NGƯỜI DÙNG: {tt_note}" if tt_note else ""}

--- TỜ TRÌNH CỦA UBND TỈNH ---
{to_trinh_text if to_trinh_text else "(Không có)"}

--- DỰ THẢO NGHỊ QUYẾT ---
{du_thao_text if du_thao_text else "(Không có)"}

--- VĂN BẢN LIÊN QUAN ---
{lien_quan_text if lien_quan_text else "(Không có)"}
"""
                            res = generate_text(prompt, use_pro=True, provider=provider_key, use_search=False)
                            st.session_state.tt_result = res
                            st.session_state.tt_ban_name = ban_select
                            st.session_state.tt_nq_display = nq_name
                            database.save_draft("Báo cáo Thẩm tra", f"Thẩm tra: {nq_name}", {"noi_dung_chinh": res, "ban": ban_select})
        with tt_c2:
            if st.button("🧹 LÀM MỚI", use_container_width=True, key="tt_clear"):
                for k in ["tt_result", "tt_ban_name", "tt_nq_display", "tt_nq_name", "tt_note"]:
                    st.session_state.pop(k, None)
                st.rerun()
    
    with tt_col2:
        if "tt_result" in st.session_state:
            st.markdown("#### 📄 Kết quả Báo cáo Thẩm tra")
            ban_display = st.session_state.get("tt_ban_name", "")
            nq_display = st.session_state.get("tt_nq_display", "")
            st.info(f"🏛️ **Ban {ban_display}** | 📋 {nq_display if nq_display else 'Dự thảo Nghị quyết'}")
            
            edited_tt = st.text_area("Nội dung báo cáo (có thể chỉnh sửa):", value=st.session_state.tt_result, height=500, key="tt_edit")
            st.session_state.tt_result = edited_tt
            
            if st.button("📄 Xuất file Word chuẩn NĐ 30", use_container_width=True, type="primary", key="tt_word"):
                out = create_tham_tra_document(
                    report_text=edited_tt,
                    ban_name=ban_display,
                    nghi_quyet_name=nq_display
                )
                if out:
                    st.download_button(
                        "⬇️ Tải Báo cáo Thẩm tra (.docx)",
                        out.getvalue(),
                        f"BaoCao_ThamTra_{ban_display.replace(' ', '_')}.docx",
                        key="tt_dl"
                    )
                else:
                    st.error("Lỗi khi tạo file Word.")
        else:
            st.info("Kết quả thẩm tra sẽ hiển thị ở đây sau khi AI hoàn tất phân tích.")
