import io
from docx import Document
import PyPDF2


def extract_text_from_pdf(file_stream):
    try:
        reader = PyPDF2.PdfReader(file_stream)
        return "\n".join(p.extract_text() or "" for p in reader.pages)
    except Exception as e:
        return f"Lỗi khi đọc PDF: {e}"


def extract_text_from_docx(file_stream):
    try:
        doc = Document(file_stream)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        return f"Lỗi khi đọc DOCX: {e}"


def create_nd30_document(content_dict):
    """Tạo file Word chuẩn Nghị định 30/2020/NĐ-CP."""
    try:
        from docx.shared import Cm, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.table import WD_TABLE_ALIGNMENT
        import re

        doc = Document()

        # 1. Page setup A4
        for section in doc.sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(2.0)

        get = content_dict.get
        cq_chu_quan = get("co_quan_chu_quan", "")
        cq_ban_hanh = get("co_quan_ban_hanh", "")
        so_ky_hieu  = get("so_ky_hieu", "")
        dia_danh    = get("dia_danh_ngay_thang", "")
        loai_vb     = get("loai_van_ban", "")
        trich_yeu   = get("trich_yeu", "")
        noi_dung    = get("noi_dung_chinh", "")
        noi_nhan    = get("noi_nhan", [])
        quyen_han   = get("quyen_han_ky", "")
        nguoi_ky    = get("nguoi_ky", "")

        def _add_run(paragraph, text, size=13, bold=False, italic=False, name='Times New Roman'):
            run = paragraph.add_run(text)
            run.font.name = name
            run.font.size = Pt(size)
            run.bold = bold
            run.italic = italic
            return run

        def _center_paragraph(cell_or_doc, is_cell=True):
            p = cell_or_doc.paragraphs[0] if is_cell else cell_or_doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return p

        # 2. Header Table
        tbl = doc.add_table(rows=2, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in tbl.rows:
            row.cells[0].width = Cm(7.0)
            row.cells[1].width = Cm(9.0)

        # Col 0 Row 0: Cơ quan chủ quản
        p = _center_paragraph(tbl.cell(0, 0))
        if cq_chu_quan:
            _add_run(p, cq_chu_quan)

        # Col 0 Row 1: Cơ quan ban hành + Số ký hiệu
        p = _center_paragraph(tbl.cell(1, 0))
        p.paragraph_format.space_after = Pt(0)
        if cq_ban_hanh:
            _add_run(p, cq_ban_hanh, bold=True)
            p_line = tbl.cell(1, 0).add_paragraph()
            p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_line.paragraph_format.space_before = Pt(0)
            p_line.paragraph_format.space_after = Pt(0)
            stars = "*" * (15 if len(cq_ban_hanh) > 20 else 10)
            _add_run(p_line, stars, size=10)

        p_so = tbl.cell(1, 0).add_paragraph()
        p_so.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if so_ky_hieu:
            _add_run(p_so, so_ky_hieu)

        # Col 1 Row 0: Quốc hiệu
        p = _center_paragraph(tbl.cell(0, 1))
        _add_run(p, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True)

        # Col 1 Row 1: Tiêu ngữ + Ngày tháng
        p = _center_paragraph(tbl.cell(1, 1))
        p.paragraph_format.space_after = Pt(0)
        _add_run(p, "Độc lập - Tự do - Hạnh phúc", size=14, bold=True)

        p_line2 = tbl.cell(1, 1).add_paragraph()
        p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_line2.paragraph_format.space_before = Pt(0)
        p_line2.paragraph_format.space_after = Pt(0)
        _add_run(p_line2, "________________________", size=12)

        p_date = tbl.cell(1, 1).add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if dia_danh:
            _add_run(p_date, dia_danh, size=14, italic=True)

        doc.add_paragraph()

        # 3. Tiêu đề văn bản
        if loai_vb:
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p_title, loai_vb, size=14, bold=True)

        if trich_yeu:
            p_trich = doc.add_paragraph()
            p_trich.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p_trich, trich_yeu, size=14, bold=(loai_vb != ""))

        doc.add_paragraph()

        # 4. Nội dung chính
        for line in noi_dung.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue

            new_p = doc.add_paragraph()
            new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            new_p.paragraph_format.first_line_indent = Cm(1.27)
            new_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            new_p.paragraph_format.line_spacing = Pt(19)
            new_p.paragraph_format.space_before = Pt(6)
            new_p.paragraph_format.space_after = Pt(6)

            is_can_cu = stripped.lower().startswith("căn cứ")

            match = re.match(r'^(\d+(?:\.\d+)*\.)\s+(.*)', stripped)
            if match:
                _add_run(new_p, match.group(1) + " ", size=14, bold=True)
                stripped = match.group(2)

            for part in re.split(r'(\*\*.*?\*\*)', stripped):
                is_bold = part.startswith('**') and part.endswith('**')
                text = part[2:-2] if is_bold else part.replace('*', '')
                if text:
                    _add_run(new_p, text, size=14, bold=is_bold, italic=is_can_cu)

        doc.add_paragraph()

        # 5. Footer Table
        tbl_f = doc.add_table(rows=1, cols=2)
        tbl_f.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl_f.rows[0].cells[0].width = Cm(7.0)
        tbl_f.rows[0].cells[1].width = Cm(9.0)

        p_nn = tbl_f.cell(0, 0).paragraphs[0]
        p_nn.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _add_run(p_nn, "Nơi nhận:", size=12, bold=True, italic=True)

        for nn in (noi_nhan or []):
            p = tbl_f.cell(0, 0).add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            _add_run(p, nn, size=11)

        p_sign = tbl_f.cell(0, 1).paragraphs[0]
        p_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if quyen_han:
            for sline in quyen_han.split('\n'):
                _add_run(p_sign, sline + '\n', size=14, bold=True)

        p_name = tbl_f.cell(0, 1).add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.paragraph_format.space_before = Pt(60)
        if nguoi_ky:
            _add_run(p_name, nguoi_ky, size=14, bold=True)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    except Exception as e:
        print(f"Lỗi khi xử lý template DOCX: {e}")
        return None
